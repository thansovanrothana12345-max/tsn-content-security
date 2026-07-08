import os
import base64
import tempfile
import io
import sqlite3
from fastapi import APIRouter, HTTPException, Response, Depends
from pydantic import BaseModel
from typing import Optional
from backend.database import get_db_connection
from backend.config import Config
from backend.routes.auth import require_role
from fpdf import FPDF
from docx import Document
from docx.shared import Inches, Pt

router = APIRouter(prefix="/api/v1/reports", tags=["Reports"])

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

class ReportGenerateRequest(BaseModel):
    case_id: int
    evidence_id: int
    sender_name: str
    sender_email: str
    sender_phone: Optional[str] = None
    sender_address: Optional[str] = None
    copyright_owner_name: str
    template_type: str = "standard"  # standard, youtube, tiktok, meta, cease_desist
    signature_base64: Optional[str] = None
    declaration_accepted: bool = False

class NoticePDF(FPDF):
    def __init__(self):
        super().__init__()
        self.set_auto_page_break(auto=True, margin=15)
        
    def header(self):
        self.set_font("Helvetica", "B", 8)
        self.set_text_color(128, 128, 128)
        self.cell(0, 10, "COPYRIGHT SECURITY - SYSTEM COMPLIANCE REPORT", ln=True, align="R")
        self.line(10, 18, 200, 18)
        self.ln(5)
        
    def footer(self):
        self.set_y(-15)
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(128, 128, 128)
        self.cell(0, 10, f"Page {self.page_no()}/{{nb}}", align="C")

def save_base64_temp_image(base64_str: str) -> Optional[str]:
    if not base64_str:
        return None
    if "," in base64_str:
        base64_str = base64_str.split(",", 1)[1]
    
    try:
        img_data = base64.b64decode(base64_str)
        temp_fd, temp_path = tempfile.mkstemp(suffix=".png")
        with os.fdopen(temp_fd, "wb") as f:
            f.write(img_data)
        return temp_path
    except Exception:
        return None

def generate_dmca_text(template_type, platform, sender_name, sender_email, sender_phone, sender_address, copyright_owner_name, original_title, infringing_url, infringing_title, uploader):
    phone_str = f"Phone: {sender_phone}\n" if sender_phone else ""
    address_str = f"Address: {sender_address}\n" if sender_address else ""
    
    # Standard DMCA Notice
    if template_type == "standard" or not template_type:
        header = f"""DMCA COPYRIGHT COMPLAINT NOTICE
To: {platform} Copyright Agent / Legal Department
Date: [Current Date]

From:
Name: {sender_name}
Email: {sender_email}
{phone_str}{address_str}
On behalf of Copyright Owner: {copyright_owner_name}

---------------------------------------------------------
RE: Notice of Copyright Infringement (DMCA Takedown Request)
---------------------------------------------------------

Dear Copyright Agent,

I am writing to you on behalf of {copyright_owner_name}, the copyright owner of the original video content described below. I have a good faith belief that the use of the material in the manner complained of is not authorized by the copyright owner, its agent, or the law.

1. IDENTIFICATION OF THE ORIGINAL WORK:
Title of Original Video: "{original_title}"
Ownership Claim: The copyright owner, {copyright_owner_name}, holds all exclusive copyrights to this original audio-visual creation.

2. IDENTIFICATION OF THE INFRINGING WORK TO BE REMOVED:
Platform: {platform}
Infringing Video URL: {infringing_url}
Infringing Video Title: "{infringing_title}"
Uploader Account / Channel: {uploader}

3. SENDER CONTACT INFORMATION:
Name: {sender_name}
Relation to Owner: Authorized Representative / Creator
Email Address: {sender_email}
{phone_str}{address_str}
"""
        declarations = f"""4. LEGAL DECLARATIONS:
* Good Faith Belief: I hereby state that I have a good faith belief that use of the copyrighted materials described above as allegedly infringing is not authorized by the copyright owner, its agent, or the law.
* Accuracy Under Penalty of Perjury: I hereby state that the information in this notification is accurate, and, under penalty of perjury, that I am the copyright owner or authorized to act on behalf of the owner of an exclusive right that is allegedly infringed.

5. ELECTRONIC SIGNATURE:
Conformed Electronic Signature: /s/ {sender_name}

Please remove the infringing material immediately. If you require further information, please contact me immediately at {sender_email}.

Thank you for your prompt cooperation.

Sincerely,
{sender_name}
"""
        return header + "\n" + declarations

    # YouTube Notice
    elif template_type == "youtube":
        return f"""YOUTUBE COPYRIGHT COMPLAINT NOTICE
To: YouTube Copyright Agent / Google LLC
Date: [Current Date]

From:
Name: {sender_name}
Email: {sender_email}
{phone_str}{address_str}
On behalf of Copyright Owner: {copyright_owner_name}

---------------------------------------------------------
RE: Notice of Copyright Infringement on YouTube
---------------------------------------------------------

Dear YouTube Copyright Team,

I am writing to request the immediate removal of infringing content uploaded to YouTube. I have a good faith belief that the use of the material in the manner complained of is not authorized by the copyright owner, its agent, or the law.

1. IDENTIFICATION OF THE ORIGINAL WORK:
Original Work Title: "{original_title}"
Ownership Claim: The copyright owner, {copyright_owner_name}, holds all exclusive copyrights to this original audio-visual creation.

2. IDENTIFICATION OF THE INFRINGING WORK TO BE REMOVED:
Infringing Video URL: {infringing_url}
Infringing Video Title: "{infringing_title}"
Uploader Account / Channel: {uploader}

3. SENDER CONTACT INFORMATION:
Name: {sender_name}
Email: {sender_email}
{phone_str}{address_str}

4. LEGAL DECLARATIONS:
* Good Faith: I have a good faith belief that use of the material in the manner complained of is not authorized by the copyright owner, its agent, or the law.
* Accuracy & Perjury: The information in this notification is accurate, and under penalty of perjury, I am the copyright owner or authorized to act on behalf of the owner of an exclusive right that is allegedly infringed.

5. ELECTRONIC SIGNATURE:
Conformed Electronic Signature: /s/ {sender_name}

Please remove the infringing video file from your platform immediately.

Sincerely,
{sender_name}
"""

    # TikTok Report
    elif template_type == "tiktok":
        return f"""TIKTOK INTELLECTUAL PROPERTY REPORT (COPYRIGHT COMPLAINT)
To: TikTok Legal Department / Bytedance
Date: [Current Date]

From:
Name: {sender_name}
Email: {sender_email}
{phone_str}{address_str}
On behalf of Copyright Owner: {copyright_owner_name}

---------------------------------------------------------
RE: TikTok Video Infringement Notification
---------------------------------------------------------

Dear TikTok IP Team,

I am reporting a copyright copyright infringement on TikTok. I have a good faith belief that the use of the material in the manner complained of is not authorized by the copyright owner, its agent, or the law.

1. IDENTIFICATION OF THE ORIGINAL WORK:
Original Work Title: "{original_title}"
Ownership Claim: Owned by {copyright_owner_name}.

2. IDENTIFICATION OF THE INFRINGING WORK TO BE REMOVED:
TikTok Video URL: {infringing_url}
Video Title/Description: "{infringing_title}"
Uploader TikTok Username: {uploader}

3. SENDER CONTACT INFORMATION:
Name: {sender_name}
Email: {sender_email}
{phone_str}{address_str}

4. LEGAL DECLARATIONS:
* Accuracy & Perjury: I represent that the information in this notification is accurate, and, under penalty of perjury, that I am the copyright owner or authorized to act on behalf of the owner of an exclusive right that is allegedly infringed.
* Good Faith: I have a good faith belief that the dispute is not authorized by the copyright owner, its agent, or the law.

5. ELECTRONIC SIGNATURE:
Conformed Electronic Signature: /s/ {sender_name}

Sincerely,
{sender_name}
"""

    # Meta Notice
    elif template_type == "meta":
        return f"""META INTELLECTUAL PROPERTY RIGHTS COMPLAINT
To: Meta Platforms Legal Department / Copyright Agent
Date: [Current Date]

From:
Name: {sender_name}
Email: {sender_email}
{phone_str}{address_str}
On behalf of Copyright Owner: {copyright_owner_name}

---------------------------------------------------------
RE: Meta Platform Copyright Infringement
---------------------------------------------------------

Dear Meta Intellectual Property Team,

I am writing to report a copyright infringement of our original audio-visual creation on Meta platforms. I have a good faith belief that the use of the material in the manner complained of is not authorized by the copyright owner, its agent, or the law.

1. IDENTIFICATION OF THE ORIGINAL WORK:
Original Work Title: "{original_title}"
Ownership Claim: Owned by {copyright_owner_name}.

2. IDENTIFICATION OF THE INFRINGING WORK TO BE REMOVED:
Meta Video/Post URL: {infringing_url}
Description: "{infringing_title}"
Page/Account Name: {uploader}

3. SENDER CONTACT INFORMATION:
Name: {sender_name}
Email: {sender_email}
{phone_str}{address_str}

4. LEGAL DECLARATIONS:
* Good Faith: I have a good faith belief that the use of the copyrighted material is not authorized by the copyright owner, its agent, or the law.
* Accuracy & Perjury: Under penalty of perjury, I declare that the information in this notice is accurate and that I am the copyright owner or authorized representative.

5. ELECTRONIC SIGNATURE:
Conformed Electronic Signature: /s/ {sender_name}

Sincerely,
{sender_name}
"""

    # Cease and Desist Demand Letter
    elif template_type == "cease_desist":
        return f"""CEASE AND DESIST DEMAND LETTER
To: {uploader}
Email / Platform Contact: Via Infringing Post/Profile
Date: [Current Date]

From:
On behalf of Copyright Owner: {copyright_owner_name}
Represented by: {sender_name} ({sender_email})

---------------------------------------------------------
RE: URGENT: NOTICE TO CEASE AND DESIST
---------------------------------------------------------

Dear {uploader},

It has come to our attention that you are hosting, displaying, or distributing unauthorized copies of our copyrighted audio-visual work: "{original_title}" on the following platform: {platform}.

Specifically, your post located at:
URL: {infringing_url}
Title/Description: "{infringing_title}"

This unauthorized use of our copyrighted work is a violation of international copyright laws.

We hereby demand that you:
1. Immediately remove the infringing content from the above URL and any other platforms where you have uploaded it.
2. Cease any further unauthorized distribution, reproduction, or performance of our copyrighted works.

If the infringing content is not removed within twenty-four (24) hours, we will pursue all available legal remedies, including submitting formal DMCA notices to the hosting platforms and domain registrars, which can result in the permanent suspension of your social media channel or page, and potentially seeking statutory damages.

You may contact the sender at {sender_email} for any queries or confirmation of removal.

Sincerely,
{sender_name}
Representative for {copyright_owner_name}
"""
    return ""

@router.get("/{case_id}")
def list_reports(case_id: int, user: dict = Depends(require_role(["Admin", "Editor", "Reviewer", "Guest"]))):
    """Lists all generated DMCA reports for a case."""
    conn = get_db_connection()
    try:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM dmca_reports WHERE case_id = ? ORDER BY created_at DESC", (case_id,))
        rows = cursor.fetchall()
        return [dict(row) for row in rows]
    finally:
        conn.close()

@router.post("/generate")
def generate_report(request: ReportGenerateRequest, user: dict = Depends(require_role(["Admin", "Editor"]))):
    """Generates and saves a DMCA report for an infringing video entry."""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    # Fetch case details
    cursor.execute("SELECT title FROM cases WHERE id = ?", (request.case_id,))
    case_row = cursor.fetchone()
    if not case_row:
        conn.close()
        raise HTTPException(status_code=404, detail="Case not found")
    case_title = case_row["title"]
    
    # Fetch evidence details
    cursor.execute("SELECT * FROM evidence WHERE id = ?", (request.evidence_id,))
    evidence_row = cursor.fetchone()
    if not evidence_row:
        conn.close()
        raise HTTPException(status_code=404, detail="Evidence not found")
        
    infringing_url = evidence_row["url"]
    infringing_title = evidence_row["title"]
    uploader = evidence_row["uploader"]
    platform = evidence_row["platform"]
    
    # Fetch corresponding original video name
    cursor.execute("SELECT filename FROM originals WHERE case_id = ? LIMIT 1", (request.case_id,))
    orig_row = cursor.fetchone()
    original_title = orig_row["filename"] if orig_row else case_title
    
    # Generate notice text
    report_text = generate_dmca_text(
        template_type=request.template_type,
        platform=platform,
        sender_name=request.sender_name,
        sender_email=request.sender_email,
        sender_phone=request.sender_phone,
        sender_address=request.sender_address,
        copyright_owner_name=request.copyright_owner_name,
        original_title=original_title,
        infringing_url=infringing_url,
        infringing_title=infringing_title,
        uploader=uploader
    )
    
    # Save to DB
    status_value = "Signed" if request.declaration_accepted else "Draft"
    cursor.execute(
        """
        INSERT INTO dmca_reports (case_id, platform, sender_name, sender_email, sender_phone, sender_address, copyright_owner_name, report_text, template_type, signature_base64, status)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            request.case_id,
            platform,
            request.sender_name,
            request.sender_email,
            request.sender_phone,
            request.sender_address,
            request.copyright_owner_name,
            report_text,
            request.template_type,
            request.signature_base64,
            status_value
        )
    )
    report_id = cursor.lastrowid
    
    # Update evidence status
    cursor.execute("UPDATE evidence SET status = 'DMCA Drafted' WHERE id = ?", (request.evidence_id,))
    
    # Insert audit trail log
    cursor.execute(
        """
        INSERT INTO audit_logs (user_id, action, entity_type, entity_id, details_json)
        VALUES (?, 'CREATE_REPORT', 'dmca_report', ?, ?)
        """,
        (user["id"], report_id, f'{{"template": "{request.template_type}", "case_id": {request.case_id}}}')
    )
    
    conn.commit()
    
    # Fetch created report
    cursor.execute("SELECT * FROM dmca_reports WHERE id = ?", (report_id,))
    new_report = cursor.fetchone()
    conn.close()
    
    return dict(new_report)

@router.get("/{report_id}/export/pdf")
def export_pdf(report_id: int, user: dict = Depends(require_role(["Admin", "Editor", "Reviewer", "Guest"]))):
    """Generates print-ready PDF containing notice metadata and digital signatures."""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    cursor.execute("SELECT * FROM dmca_reports WHERE id = ?", (report_id,))
    report_row = cursor.fetchone()
    if not report_row:
        conn.close()
        raise HTTPException(status_code=404, detail="Report not found")
    report = dict(report_row)
    
    cursor.execute("SELECT title FROM cases WHERE id = ?", (report["case_id"],))
    case_title = cursor.fetchone()["title"]
    
    # Find matching evidence screenshot
    # Check if there is an evidence row for the platform and case
    cursor.execute(
        "SELECT platform, url, uploader, similarity_score, screenshot_path FROM evidence WHERE case_id = ? AND platform = ? LIMIT 1",
        (report["case_id"], report["platform"])
    )
    ev_row = cursor.fetchone()
    evidence = dict(ev_row) if ev_row else {
        "platform": report["platform"],
        "url": "N/A",
        "uploader": "N/A",
        "similarity_score": 0.0,
        "screenshot_path": None
    }
    
    cursor.execute("SELECT filename FROM originals WHERE case_id = ? LIMIT 1", (report["case_id"],))
    orig_row = cursor.fetchone()
    original_title = orig_row["filename"] if orig_row else case_title
    
    conn.close()
    
    # Resolve screenshot path
    screenshot_full_path = None
    if evidence.get("screenshot_path"):
        filename = os.path.basename(evidence["screenshot_path"])
        screenshot_full_path = os.path.join(PROJECT_ROOT, "storage", "evidence", filename)
    
    # Create PDF
    pdf = NoticePDF()
    pdf.add_page()
    
    pdf.set_font("Helvetica", "B", 18)
    pdf.set_text_color(33, 33, 33)
    pdf.cell(0, 10, "Copyright Infringement Notice Report", ln=True, align="L")
    pdf.ln(5)
    
    # Details Grid
    pdf.set_font("Helvetica", "B", 10)
    pdf.set_text_color(64, 64, 64)
    
    details = [
        ("Case Folder:", case_title),
        ("Target Platform:", evidence["platform"]),
        ("Infringing URL:", evidence["url"]),
        ("Uploader Channel:", evidence["uploader"]),
        ("Original Title Reference:", original_title),
        ("Visual Match Score:", f"{evidence['similarity_score'] * 100:.1f}% Similarity Match")
    ]
    
    for label, val in details:
        pdf.set_font("Helvetica", "B", 10)
        pdf.cell(45, 8, label, border=0)
        pdf.set_font("Helvetica", "", 10)
        if label == "Infringing URL:" and val != "N/A":
            pdf.set_text_color(30, 80, 200)
            pdf.cell(0, 8, val, ln=True)
            pdf.set_text_color(64, 64, 64)
        else:
            pdf.cell(0, 8, val, ln=True)
            
    pdf.ln(5)
    
    # Visual proof screenshot
    if screenshot_full_path and os.path.exists(screenshot_full_path):
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 10, "Visual Evidence Reference", ln=True)
        pdf.ln(2)
        try:
            pdf.image(screenshot_full_path, w=180)
        except Exception:
            pdf.cell(0, 8, "[Error loading screenshot image file]", ln=True)
        pdf.ln(10)
        
    # Legal notice text
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 10, "Drafted Legal Notice Text", ln=True)
    pdf.ln(2)
    pdf.set_font("Courier", "", 9)
    pdf.multi_cell(0, 5, report["report_text"], border=1, fill=False)
    pdf.ln(8)
    
    # Signature
    if report["signature_base64"]:
        sig_img_path = save_base64_temp_image(report["signature_base64"])
        if sig_img_path and os.path.exists(sig_img_path):
            pdf.set_font("Helvetica", "B", 12)
            pdf.cell(0, 10, "Authorized Electronic Signature", ln=True)
            pdf.ln(2)
            try:
                pdf.image(sig_img_path, w=50)
            except Exception:
                pass
            pdf.ln(5)
            os.remove(sig_img_path)
            
    pdf_bytes = pdf.output()
    
    return Response(
        content=bytes(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=DMCA_Notice_Report_{report_id}.pdf"}
    )

@router.get("/{report_id}/export/docx")
def export_docx(report_id: int, user: dict = Depends(require_role(["Admin", "Editor", "Reviewer", "Guest"]))):
    """Generates Word DOCX containing notice metadata and digital signatures."""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    cursor.execute("SELECT * FROM dmca_reports WHERE id = ?", (report_id,))
    report_row = cursor.fetchone()
    if not report_row:
        conn.close()
        raise HTTPException(status_code=404, detail="Report not found")
    report = dict(report_row)
    
    cursor.execute("SELECT title FROM cases WHERE id = ?", (report["case_id"],))
    case_title = cursor.fetchone()["title"]
    
    cursor.execute(
        "SELECT platform, url, uploader, similarity_score, screenshot_path FROM evidence WHERE case_id = ? AND platform = ? LIMIT 1",
        (report["case_id"], report["platform"])
    )
    ev_row = cursor.fetchone()
    evidence = dict(ev_row) if ev_row else {
        "platform": report["platform"],
        "url": "N/A",
        "uploader": "N/A",
        "similarity_score": 0.0,
        "screenshot_path": None
    }
    
    cursor.execute("SELECT filename FROM originals WHERE case_id = ? LIMIT 1", (report["case_id"],))
    orig_row = cursor.fetchone()
    original_title = orig_row["filename"] if orig_row else case_title
    
    conn.close()
    
    screenshot_full_path = None
    if evidence.get("screenshot_path"):
        filename = os.path.basename(evidence["screenshot_path"])
        screenshot_full_path = os.path.join(PROJECT_ROOT, "storage", "evidence", filename)
        
    doc = Document()
    
    # Document Style Title
    title = doc.add_paragraph()
    title_run = title.add_run("Copyright Infringement Notice Report")
    title_run.font.name = "Arial"
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    
    doc.add_heading("Case & Verification Details", level=1)
    
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    details_map = [
        ("Case Folder", case_title),
        ("Target Platform", evidence["platform"]),
        ("Infringing URL", evidence["url"]),
        ("Uploader Channel", evidence["uploader"]),
        ("Match Score", f"{evidence['similarity_score'] * 100:.1f}% Similarity Match")
    ]
    
    for i, (k, v) in enumerate(details_map):
        row_cells = table.rows[i].cells
        row_cells[0].text = k
        row_cells[1].text = v
        
    doc.add_paragraph()
    
    # Screenshot Proof
    if screenshot_full_path and os.path.exists(screenshot_full_path):
        doc.add_heading("Visual Evidence Reference", level=1)
        doc.add_paragraph()
        try:
            doc.add_picture(screenshot_full_path, width=Inches(6.0))
        except Exception:
            doc.add_paragraph("[Error loading screenshot image]")
        doc.add_paragraph()
        
    # Legal notice text
    doc.add_heading("Drafted Legal Notice Text", level=1)
    notice_p = doc.add_paragraph()
    notice_run = notice_p.add_run(report["report_text"])
    notice_run.font.name = "Courier New"
    notice_run.font.size = Pt(10)
    
    # Signature
    if report["signature_base64"]:
        sig_img_path = save_base64_temp_image(report["signature_base64"])
        if sig_img_path and os.path.exists(sig_img_path):
            doc.add_heading("Authorized Signature", level=1)
            doc.add_paragraph()
            try:
                doc.add_picture(sig_img_path, width=Inches(2.0))
            except Exception:
                pass
            os.remove(sig_img_path)
            
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return Response(
        content=file_stream.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=DMCA_Notice_Report_{report_id}.docx"}
    )

@router.delete("/{report_id}")
def delete_report(report_id: int, user: dict = Depends(require_role(["Admin"]))):
    """Deletes a generated report."""
    conn = get_db_connection()
    try:
        cursor = conn.cursor()
        cursor.execute("SELECT id FROM dmca_reports WHERE id = ?", (report_id,))
        if not cursor.fetchone():
            raise HTTPException(status_code=404, detail="Report not found")
            
        cursor.execute("DELETE FROM dmca_reports WHERE id = ?", (report_id,))
        conn.commit()
    finally:
        conn.close()
    return {"message": f"Report {report_id} deleted successfully."}


class LogTakedownRequest(BaseModel):
    evidence_id: int
    recipient_platform: str
    action_taken: str
    status: Optional[str] = "Draft"
    legal_signee: Optional[str] = None

class UpdateTakedownStatusRequest(BaseModel):
    status: str

@router.post("/takedown/logs", status_code=201)
def log_takedown_action_api(
    request: LogTakedownRequest,
    user: dict = Depends(require_role(["Admin", "Editor"]))
):
    """Logs an enforcement takedown notice audit log entry."""
    from backend.services.copyright_enforcement import CopyrightEnforcementService
    try:
        log_id = CopyrightEnforcementService.log_takedown(
            evidence_id=request.evidence_id,
            recipient_platform=request.recipient_platform,
            action_taken=request.action_taken,
            status=request.status,
            legal_signee=request.legal_signee
        )
        return {"log_id": log_id, "message": "Copyright takedown enforcement action successfully logged."}
    except ValueError as val_err:
        raise HTTPException(status_code=404, detail=str(val_err))
    except Exception as err:
        raise HTTPException(status_code=500, detail=str(err))

@router.get("/takedown/logs/{evidence_id}")
def get_takedown_logs_api(
    evidence_id: int,
    user: dict = Depends(require_role(["Admin", "Editor", "Reviewer", "Guest"]))
):
    """Fetches the complete enforcement history logs for a specific evidence file."""
    from backend.services.copyright_enforcement import CopyrightEnforcementService
    try:
        history = CopyrightEnforcementService.get_takedown_history(evidence_id)
        return history
    except Exception as err:
        raise HTTPException(status_code=500, detail=str(err))

@router.post("/takedown/logs/{log_id}/status")
def update_takedown_status_api(
    log_id: int,
    request: UpdateTakedownStatusRequest,
    user: dict = Depends(require_role(["Admin", "Editor"]))
):
    """Transitions the status of a logged copyright takedown notice."""
    from backend.services.copyright_enforcement import CopyrightEnforcementService
    try:
        success = CopyrightEnforcementService.update_takedown_status(log_id, request.status)
        if not success:
            raise HTTPException(status_code=400, detail="Failed to update takedown log status.")
        return {"message": "Takedown status updated successfully."}
    except ValueError as val_err:
        raise HTTPException(status_code=404, detail=str(val_err))
    except Exception as err:
        raise HTTPException(status_code=500, detail=str(err))


@router.get("/compliance")
def get_compliance_report(
    user: dict = Depends(require_role(["Admin"]))
):
    """Generates a signed enterprise security and audit compliance summary report."""
    import hashlib
    import json
    from datetime import datetime
    from backend.services.audit_verifier import AuditVerifier
    
    # 1. Verify log chain integrity
    chain_ok, chain_msg = AuditVerifier.verify_chain()
    
    # 2. Fetch users summary
    conn = get_db_connection()
    try:
        cursor = conn.cursor()
        cursor.execute("SELECT id, username, email, role, created_at FROM users;")
        users = [dict(u) for u in cursor.fetchall()]
    finally:
        conn.close()
        
    # 3. Create raw payload
    payload = {
        "compliance_status": "COMPLIANT" if chain_ok else "NON_COMPLIANT",
        "audit_logs_integrity": {
            "valid": chain_ok,
            "message": chain_msg
        },
        "active_users_count": len(users),
        "users": users,
        "hardened_headers": {
            "X-Content-Type-Options": "nosniff",
            "X-Frame-Options": "DENY",
            "Content-Security-Policy": "Active"
        },
        "timestamp": datetime.utcnow().isoformat()
    }
    
    # 4. Generate digital cryptographic signature
    payload_str = json.dumps(payload, sort_keys=True)
    signature = hashlib.sha256(f"{payload_str}|{Config.SECRET_KEY}".encode('utf-8')).hexdigest()
    
    return {
        "report_payload": payload,
        "cryptographic_signature": signature
    }
